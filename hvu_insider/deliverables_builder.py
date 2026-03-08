"""
Deliverable 2: Deliverables by Team Generator
-----------------------------------------------
Produces the operational playbook .docx — every deliverable, every owner,
every deadline, organized by team section.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Brand Colors ──────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GRAY = RGBColor(0x66, 0x66, 0x66)


def _run(paragraph, text, size=Pt(10.5), color=None, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return run


def _section_header(doc, text, level=1):
    """Add a navy bold section header."""
    sizes = {1: Pt(14), 2: Pt(12), 3: Pt(11)}
    p = doc.add_paragraph()
    _run(p, text, size=sizes.get(level, Pt(11)), color=NAVY, bold=True)
    return p


def _task_item(doc, task: str, owner: str, due: str = "", deps: str = ""):
    """Add a checkbox task item with owner/due metadata."""
    p = doc.add_paragraph()
    _run(p, f"\u2610  {task}", size=Pt(10.5))

    meta_parts = [f"Owner: {owner}"]
    if due:
        meta_parts.append(f"Due: {due}")
    if deps:
        meta_parts.append(f"Deps: {deps}")

    p2 = doc.add_paragraph()
    _run(p2, "    " + " | ".join(meta_parts), size=Pt(9), color=GRAY, italic=True)


def _sub_header(doc, text, side_label: str = ""):
    """Performance/Human side sub-header."""
    p = doc.add_paragraph()
    _run(p, text, size=Pt(12), color=NAVY, bold=True)
    if side_label:
        _run(p, f"  [{side_label}]", size=Pt(9), color=GOLD, italic=True)


def build_deliverables_doc(
    position: str,
    year: int,
    players: list[dict],
    drop_date: str,
    output_dir: Path,
) -> Path:
    """Build the full deliverables-by-team .docx."""
    doc = Document()
    n = len(players)
    pos = position.upper()
    player_names = [p["name"] for p in players]
    names_str = ", ".join(player_names)

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1))

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    # ── Header Block ─────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"HVU INSIDER | {year} {pos} INTEL DROP | DELIVERABLES BY TEAM",
         size=Pt(14), color=NAVY, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, f"POSITION GROUP: {pos} | {n} Targets", size=Pt(11), color=NAVY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, names_str, size=Pt(10), color=GRAY)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p4, "Every deliverable. Every owner. Every deadline. Two sides of every recruit.",
         size=Pt(10), color=GRAY, italic=True)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p5, "REUSABLE TEMPLATE — Same format for every future position group drop",
         size=Pt(9), color=GOLD, bold=True)

    doc.add_paragraph()

    # ── Two Sides Framework ──────────────────────────────────────────────
    _section_header(doc, "THE TWO SIDES FRAMEWORK")
    p = doc.add_paragraph()
    _run(p, "1. Performance Side", size=Pt(10.5), bold=True)
    _run(p, " — Film, stats, scouting, scheme fit, recruiting intel\n")
    _run(p, "2. Human Side", size=Pt(10.5), bold=True)
    _run(p, " — Personality, social media, interests, lifestyle, off-field identity")

    # ── Fan Scouting Funnel ──────────────────────────────────────────────
    _section_header(doc, "FAN SCOUTING FUNNEL")
    for level in [
        "TOP OF FUNNEL — Social (comments, likes, shares, remixes)",
        "MID-FUNNEL — Insider Platform (polls, quizzes, pick'ems, keywords, challenges)",
        "BOTTOM OF FUNNEL — Message boards with Mike V",
    ]:
        p = doc.add_paragraph()
        _run(p, f"\u25AA  {level}", size=Pt(10))

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: MIKE V
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "SECTION 1: MIKE V — YOUR DELIVERABLES")

    _sub_header(doc, "1A. Performance Side — Intel & Scouting Content", "Performance")
    _task_item(doc, f"Record full {pos} video breakdown (10-15 min)", "Mike V",
               "Day Before", "Film clips ready")
    _task_item(doc, f"Provide written scouting notes for all {n} players", "Mike V",
               "Day Before")
    _task_item(doc, "Record intern scout evaluator segment (60-90 sec each)",
               "Mike V + Evaluator Interns", "Day Before")

    _sub_header(doc, "1B. Human Side — Recruit Personality Intel", "Human")
    _task_item(doc, "Research each player's public social media (IG, TikTok)",
               "Mike V", "Day Before")
    _task_item(doc, "Ask 3 quick questions if direct contact: pregame song, go-to meal, what to see at PSU",
               "Mike V", "Day Before")
    _task_item(doc, "Provide color commentary on each player's personality",
               "Mike V", "Day Before")

    _sub_header(doc, "1C. Social Posts — Mike V's Personal Accounts (9 posts)")
    _section_header(doc, "Performance Side Posts (7)", level=3)
    perf_posts = [
        ("Day Before AM (Radio)", "Tease board, mention context narrative, set drop time"),
        ("Day Before 12 PM", "Context narrative post (loss/flip/need)"),
        ("Drop Day 7 AM", 'Redacted board (████), competition angle, "full intel at 10"'),
        ("Drop Day 6 PM", 'Top target spotlight with stats, "the rest lives on the board"'),
        ("Drop +1, 7 AM", "Hot take on position room, poll CTA (+10 XP)"),
        ("Drop +1, 8 AM", "Second target spotlight, OV angle"),
        ("Drop +2, 6 PM", "Weekend debate prompt, forum push"),
    ]
    for time, desc in perf_posts:
        _task_item(doc, f"{desc}", "Mike V", time)

    _section_header(doc, "Human Side Posts (2)", level=3)
    _task_item(doc, "Personality/playlist/TikTok hook -> platform poll", "Mike V", "Drop Day 3 PM")
    _task_item(doc, "Board thread tease -> forum push", "Mike V", "Drop +1, 3 PM")

    _sub_header(doc, "1D. Platform / Forum (6 items)")
    _section_header(doc, "Performance Side Threads", level=3)
    _task_item(doc, f'Seed Forum Thread #1: "{year} {pos} Board — Your Reactions"',
               "Mike V", "Drop Day 10 AM")
    _task_item(doc, "Seed Forum Thread #2: Context narrative thread",
               "Mike V", "Drop Day 6 PM")
    _task_item(doc, f'Seed Forum Thread #3: "Weekend Debate: Most Realistic {pos} Commit?"',
               "Mike V", "Drop +2")
    _section_header(doc, "Human Side Threads", level=3)
    _task_item(doc, f'Seed Forum Thread #4: "Get to Know the {year} {pos} Targets — Off the Field"',
               "Mike V", "Drop Day")
    _task_item(doc, "Respond to 5+ fan posts personally across Drop Day and +1",
               "Mike V", "Drop Day - Drop +1")
    _task_item(doc, "Record 60-sec Scout Watch Recap to camera",
               "Mike V", "Drop +2")

    _sub_header(doc, "1E. Landon Collab + Swarm (5 items)")
    _task_item(doc, "Identify 1 unreleased intel nugget for Landon collab", "Mike V", "Day Before")
    _task_item(doc, "Record Landon collab phone call clip (20-30 sec)",
               "Mike V + Landon", "Day Before")
    _task_item(doc, "IG Post — Landon collab response", "Mike V", "Drop Day 10:15 AM")
    _task_item(doc, "Text all 20 swarm accounts with pre-written copy",
               "Mike V", "Drop Day before 10 AM")
    _task_item(doc, "Nudge swarm accounts that haven't posted by Drop +1", "Mike V", "Drop +1")

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: INTERN SCOUT EVALUATORS
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "SECTION 2: INSIDER INTERN SCOUT EVALUATORS")
    p = doc.add_paragraph()
    _run(p, "12-hour turnaround max.", size=Pt(10), color=GOLD, bold=True)

    _sub_header(doc, "2A. Performance Side — Written Content (5 items)", "Performance")
    _task_item(doc, f"Full written article: {pos} Target Board scouting intel (Brady Jordan voice)",
               "Writer/Reporter (Intern)", "Drop Day 10 AM")
    _task_item(doc, "Top target standalone deep dive", "Writer/Reporter (Intern)", "Drop Day")
    _task_item(doc, "Intern scout evaluator write-ups (200-300 words each)",
               "Evaluator Interns", "Drop Day")
    _task_item(doc, "Context narrative post-mortem article", "Writer/Reporter (Intern)", "Drop +1")
    _task_item(doc, "Scout Watch Recap (written companion to 60-sec video)",
               "Writer/Reporter (Intern)", "Drop +2")

    _sub_header(doc, "2B. Human Side — Recruit Personality Content (4 items)", "Human")
    _task_item(doc, "Research all targets' public IG + TikTok. Document personality moments.",
               "Intern Team", "Day Before")
    _task_item(doc, '"Get to Know" profile for each target (150-200 words)',
               "Intern Team", "Drop Day")
    _task_item(doc, f'"Which {pos} Are You?" personality quiz concept',
               "Community Mgr", "Drop Day")
    _task_item(doc, "Pull 3-5 short clips from recruit TikToks/IG Reels (public only)",
               "Intern Team", "Day Before")

    _sub_header(doc, "2C. Video Production (8 items)")
    _task_item(doc, "Edit Mike V's full breakdown video (10-15 min)", "Video Producer (Intern)", "Drop Day 9 AM")
    _task_item(doc, "Edit intern scout evaluator clips (60-90 sec each)", "Video Producer (Intern)", "Drop Day")
    _task_item(doc, "Cut 15-sec Board Tease Reel", "Clip Editor (Intern)", "Day Before")
    _task_item(doc, "Cut top target highlight (30 sec)", "Clip Editor (Intern)", "Drop Day")
    _task_item(doc, "Cut second target feature clip", "Clip Editor (Intern)", "Drop +1")
    _task_item(doc, "Cut regional/storyline target feature", "Clip Editor (Intern)", "Drop +1")
    _task_item(doc, "Cut Mike V Scout Watch Recap (60 sec)", "Video Producer (Intern)", "Drop +2")
    _task_item(doc, "Cut context narrative clip", "Clip Editor (Intern)", "Drop Day")

    _sub_header(doc, "2D. Social Execution (3 items)")
    _task_item(doc, "Schedule all HVU Insider account posts per content tracker",
               "Leah (Social Director)", "Day Before")
    _task_item(doc, "Monitor and respond to EVERY comment/DM on all drop posts",
               "Intern Team + Leah", "Drop Day - Drop +2")
    _task_item(doc, "Create IG Story Carousel copy (5 cards: Intro -> Silhouette -> Poll -> Stats -> CTA)",
               "Graphics Lead", "Day Before")

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: HVU INSIDER ACCOUNT POSTS
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "SECTION 3: HVU INSIDER ACCOUNT POSTS")
    p = doc.add_paragraph()
    _run(p, "Posts from @HVU_Insider (X) and @hvu_insider (IG). Bryan approves all copy.",
         size=Pt(10), color=GRAY, italic=True)

    _sub_header(doc, "3A. X (Twitter) — @HVU_Insider (10 posts)")
    _section_header(doc, "Performance Side (6)", level=3)
    x_perf = [
        ("Drop Day 10 AM", "Launch post — article link + hero graphic"),
        ("Drop Day 12 PM", "Top target feature card + stats"),
        ("Drop Day 3 PM", "Context narrative post"),
        ("Drop +1, 10 AM", "Second target feature"),
        ("Drop +2, 10 AM", "Recap post"),
        ("Saturday 12 PM", "Saturday recap"),
    ]
    for time, desc in x_perf:
        _task_item(doc, desc, "Leah/Interns", time)

    _section_header(doc, "Human Side (2)", level=3)
    _task_item(doc, "Personality hook", "Leah/Interns", "Drop Day 2 PM")
    _task_item(doc, "TikTok tease", "Leah/Interns", "Drop +1, 2 PM")

    _section_header(doc, "Engagement (2)", level=3)
    _task_item(doc, "Sweepstakes reminder", "Leah/Interns", "Drop +1, 3 PM")
    _task_item(doc, "Keyword hunt", "Leah/Interns", "Drop +2, 3 PM")

    _sub_header(doc, "3B. Instagram — @hvu_insider (12 posts)")
    _section_header(doc, "Performance Feed Posts (5)", level=3)
    ig_perf = [
        ("Drop Day 10 AM", "Redacted board"),
        ("Drop Day 3 PM", "Context narrative"),
        ("Drop +1, 10 AM", "Second target"),
        ("Drop +1, 12 PM", "Reels"),
        ("Drop +2, 10 AM", "Recap"),
    ]
    for time, desc in ig_perf:
        _task_item(doc, desc, "Leah/Interns", time)

    _section_header(doc, "Human Feed (1) + Stories (4) + Engagement (2)", level=3)
    _task_item(doc, "TikTok reaction reel", "Leah/Interns", "Drop +1")
    _task_item(doc, "Story Carousel", "Leah/Interns", "Drop Day 10 AM")
    _task_item(doc, "Top target highlight story", "Leah/Interns", "Drop Day 12 PM")
    _task_item(doc, "Personality poll story", "Leah/Interns", "Drop Day 2 PM")
    _task_item(doc, "TikTok reaction story", "Leah/Interns", "Drop +1")
    _task_item(doc, "Sweepstakes story", "Leah/Interns", "Drop +1, 3 PM")
    _task_item(doc, "Keyword story", "Leah/Interns", "Drop +2, 3 PM")

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: ODYSSEY GRAPHICS & VIDEO
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "SECTION 4: ODYSSEY GRAPHICS & VIDEO EDITING")
    p = doc.add_paragraph()
    _run(p, "Ronny + Jonah own design. Rylan + editing team own video. Bryan approves all.",
         size=Pt(9), color=GRAY, italic=True)
    p2 = doc.add_paragraph()
    _run(p2, "COLORS: Penn State Navy #041E42 | White #FFFFFF | Accent: Light Blue #A2D2FF",
         size=Pt(9), color=NAVY, bold=True)

    _sub_header(doc, "4A. Graphics — Priority (Before Drop Day) (4)")
    _task_item(doc, f"HERO: Mike V + {pos} Board (1080x1350 + 1600x900)", "Ronny/Jonah", "Day Before")
    _task_item(doc, "Redacted Board Graphic (1080x1080)", "Ronny/Jonah", "Day Before")
    _task_item(doc, "Landon x Mike V Split Design (1080x1350 + 1080x1920)", "Ronny/Jonah", "Day Before")
    _task_item(doc, "IG Story Carousel (1080x1920 x5 cards)", "Ronny/Jonah", "Day Before")

    _sub_header(doc, "4B. Graphics — Drop Day (4)")
    _task_item(doc, "Top Target Feature Card (1080x1350)", "Ronny/Jonah", "Drop Day AM")
    _task_item(doc, "Context Narrative Graphic (1080x1080)", "Ronny/Jonah", "Drop Day AM")
    _task_item(doc, "Mike V Quote Card (1080x1080 + 1080x1920)", "Ronny/Jonah", "Drop Day AM")
    _task_item(doc, "Personality Poll Graphic (1080x1920)", "Ronny/Jonah", "Drop Day AM")

    _sub_header(doc, "4C. Graphics — Drop +1/+2/Saturday (6)")
    _task_item(doc, "Second Target Card (1080x1350)", "Ronny/Jonah", "Drop Day PM")
    _task_item(doc, "Regional/Storyline Target Card (1080x1350)", "Ronny/Jonah", "Drop Day PM")
    _task_item(doc, "Sweepstakes Reminder Card (1080x1080 + 1080x1920)", "Ronny/Jonah", "Drop +1")
    _task_item(doc, "Keyword Hunt Graphic (1080x1920)", "Ronny/Jonah", "Drop +1")
    _task_item(doc, "Quiz Promo Card (1080x1080)", "Ronny/Jonah", "Drop +1")
    _task_item(doc, "Recap Graphic (1080x1080)", "Ronny/Jonah", "Drop +2")

    _sub_header(doc, "4D. Video Post-Production (3)")
    _task_item(doc, "Polish Mike V breakdown for paid media", "Rylan/Editing", "Drop Day")
    _task_item(doc, "Create 3 paid media ad cuts: 15-sec hook, 30-sec teaser, 10-sec Stories",
               "Rylan/Editing", "Drop +1")
    _task_item(doc, "Edit intern evaluator clips with branded lower thirds",
               "Rylan/Editing", "Drop Day")

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: DROPT PLATFORM
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "SECTION 5: DROPT PLATFORM DELIVERABLES")

    _sub_header(doc, "5A. Content Upload")
    _section_header(doc, "Performance Side (6)", level=3)
    for item in [
        "Main article", "Deep dive", "Evaluator clips",
        "Standalone articles", "Scout watch recap", "Context post-mortem",
    ]:
        _task_item(doc, item, "Platform Team", "Drop Day")
    _section_header(doc, "Human Side (2)", level=3)
    _task_item(doc, '"Get to Know" profiles', "Platform Team", "Drop Day")
    _task_item(doc, "TikTok reaction content", "Platform Team", "Drop +1")
    _section_header(doc, "Shared (2)", level=3)
    _task_item(doc, "Stories carousel", "Platform Team", "Drop Day")
    _task_item(doc, "Sponsored story card", "Platform Team", "Drop Day")

    _sub_header(doc, "5B. Games — All Uploaded Before Drop Day 10 AM")
    _section_header(doc, "Performance Side", level=3)
    _task_item(doc, "6 Trivia questions (+5 XP each)", "Platform Team", "Day Before")
    _task_item(doc, "3 Pick'Ems (+5 XP each)", "Platform Team", "Day Before")
    _task_item(doc, "2-3 Polls (+10 XP each)", "Platform Team", "Day Before")
    _task_item(doc, "1 Challenge (+1,000 XP + 3 sweeps)", "Platform Team", "Day Before")
    _task_item(doc, "1 Keyword (+20 XP)", "Platform Team", "Day Before")
    _section_header(doc, "Human Side", level=3)
    _task_item(doc, "1 Personality Poll (+10 XP)", "Platform Team", "Day Before")
    _task_item(doc, "1 Match Quiz (+10-40 XP)", "Platform Team", "Day Before")
    _task_item(doc, "1 Lifestyle Poll (+10 XP)", "Platform Team", "Day Before")
    _section_header(doc, "Sponsor / Zero-Party", level=3)
    _task_item(doc, "2 data capture questions (+5 XP each)", "Platform Team", "Day Before")

    _sub_header(doc, "5C. Push Notifications (2)")
    _task_item(doc, "Launch push notification", "Platform Team", "Drop Day 10 AM")
    _task_item(doc, "Keyword reminder push", "Platform Team", "Drop +2")

    _sub_header(doc, "5D. Platform QA (2)")
    _task_item(doc, "Points verification — all games awarding correct XP", "Platform Team", "Day Before PM")
    _task_item(doc, "Forum visibility check — threads live and indexed", "Platform Team", "Drop Day 10 AM")

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: FEATURE ARTICLES
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "SECTION 6: FEATURE ARTICLES")

    _sub_header(doc, "6A. Performance Side (5)", "Performance")
    for item in [
        "Primary intel drop article",
        "Top target deep dive",
        "Context post-mortem",
        "Scout watch recap",
        "Intern evaluator takes",
    ]:
        _task_item(doc, item, "Writer/Reporter (Intern)", "Per timeline")

    _sub_header(doc, "6B. Human Side (2)", "Human")
    _task_item(doc, "Combined personality profiles", "Writer/Reporter (Intern)", "Drop +1")
    _task_item(doc, '"We Asked — They Answered" (if direct contact obtained)',
               "Writer/Reporter (Intern)", "Drop +2")

    # ══════════════════════════════════════════════════════════════════════
    # MASTER TIMELINE
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "MASTER TIMELINE")
    timeline = [
        "Day Before -> Drop Day -> Drop +1 -> Drop +2 -> Saturday",
        "Same cadence every time.",
    ]
    for line in timeline:
        p = doc.add_paragraph()
        _run(p, line, size=Pt(10.5), color=NAVY, bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, '"The players change. The machine does not."',
         size=Pt(11), color=GOLD, bold=True, italic=True)

    # ── Save ─────────────────────────────────────────────────────────────
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"HVU_{position.upper()}_Deliverables_{year}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return filepath
