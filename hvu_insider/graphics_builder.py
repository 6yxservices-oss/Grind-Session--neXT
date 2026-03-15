"""
Graphics Brief Generator for Nano Banana
------------------------------------------
Since Nano Banana is browser-based (no API), this module produces:

1. A structured graphics brief (.docx) — every graphic needed, dimensions,
   text overlays, and which template to use in Nano Banana
2. A CSV data feed — player names, stats, copy per graphic slot, ready
   for Nano Banana's template field swaps
3. A JSON manifest — machine-readable version for any future API integration

Bryan opens Nano Banana, loads the template, swaps the fields using
the brief/CSV as the source of truth. The agent does 100% of the
thinking — Bryan just executes the swaps.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x1B, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _run(paragraph, text, size=Pt(10.5), color=None, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return run


def _set_cell_shading(cell, hex_color: str):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


# ── Template Definitions ─────────────────────────────────────────────────────
# Each template maps to a Nano Banana template. The "fields" are what Bryan
# swaps in the browser. The agent fills in the values per player/graphic.

GRAPHIC_TEMPLATES = [
    # ── Priority (Before Drop Day) ──────────────────────────────────────
    {
        "id": "hero_board",
        "name": "HERO: Mike V + {position} Board",
        "template_name": "HVU_Hero_Board",
        "dimensions": ["1080x1350 (Feed)", "1600x900 (Twitter)"],
        "phase": "Before Drop Day",
        "fields": {
            "position_group": "{position}",
            "year": "{year}",
            "player_count": "{n_players}",
            "player_names": "{all_names}",
        },
        "notes": "Dark navy background. Mike V photo left. Board graphic right. Premium feel.",
    },
    {
        "id": "redacted_board",
        "name": "Redacted Board",
        "template_name": "HVU_Redacted_Board",
        "dimensions": ["1080x1080 (Square)"],
        "phase": "Before Drop Day",
        "fields": {
            "position_group": "{position}",
            "year": "{year}",
            "redacted_names": "{redacted}",
        },
        "notes": "Names blacked out (████). Tease format. 'Full intel at 10 AM' text.",
    },
    {
        "id": "landon_split",
        "name": "Landon x Mike V Split Design",
        "template_name": "HVU_Landon_Split",
        "dimensions": ["1080x1350 (Feed)", "1080x1920 (Story)"],
        "phase": "Before Drop Day",
        "fields": {
            "position_group": "{position}",
            "collab_text": "INSIDER INTEL DROP",
        },
        "notes": "Split design — Mike V left, Landon right. Navy/white.",
    },
    {
        "id": "story_carousel",
        "name": "IG Story Carousel (5 cards)",
        "template_name": "HVU_Story_Carousel",
        "dimensions": ["1080x1920 x5"],
        "phase": "Before Drop Day",
        "per_card": [
            "Card 1: INTRO — '{year} {position} INTEL DROP' + 'Swipe to scout'",
            "Card 2: SILHOUETTE — Top target silhouette + '???'",
            "Card 3: POLL — 'Which {position} fits PSU best?' + player names as options",
            "Card 4: STATS — Top target key stats overlay",
            "Card 5: CTA — 'Full intel live on the platform' + link sticker",
        ],
    },
    # ── Drop Day ────────────────────────────────────────────────────────
    {
        "id": "top_target_card",
        "name": "Top Target Feature Card",
        "template_name": "HVU_Player_Feature",
        "dimensions": ["1080x1350 (Feed)"],
        "phase": "Drop Day",
        "per_player": True,
        "player_index": 0,
        "fields": {
            "player_name": "{player_name}",
            "school": "{school}",
            "measurables": "{height} | {weight} lbs",
            "position": "{position}",
            "headline_stat": "{top_stat}",
        },
        "notes": "Feature card for #1 target. Photo placeholder. Key stat callout.",
    },
    {
        "id": "context_narrative",
        "name": "Context Narrative Graphic",
        "template_name": "HVU_Context_Narrative",
        "dimensions": ["1080x1080 (Square)"],
        "phase": "Drop Day",
        "fields": {
            "headline": "{context_headline}",
            "subtext": "{context_subtext}",
        },
        "notes": "The 'why now' graphic. Dark background, white text, urgency feel.",
    },
    {
        "id": "mike_v_quote",
        "name": "Mike V Quote Card",
        "template_name": "HVU_Quote_Card",
        "dimensions": ["1080x1080 (Square)", "1080x1920 (Story)"],
        "phase": "Drop Day",
        "fields": {
            "quote_text": "[Mike V quote — to be provided]",
            "attribution": "— Mike V, HVU Insider",
        },
        "notes": "Pull quote from Mike V's scouting notes. Navy background.",
    },
    {
        "id": "personality_poll",
        "name": "Personality Poll Graphic",
        "template_name": "HVU_Poll",
        "dimensions": ["1080x1920 (Story)"],
        "phase": "Drop Day",
        "fields": {
            "question": "Which {position} target's vibe matches yours?",
            "options": "{all_names}",
        },
        "notes": "IG Story poll format. Fun, personality-focused.",
    },
    # ── Drop +1/+2/Saturday ─────────────────────────────────────────────
    {
        "id": "second_target_card",
        "name": "Second Target Feature Card",
        "template_name": "HVU_Player_Feature",
        "dimensions": ["1080x1350 (Feed)"],
        "phase": "Drop +1",
        "per_player": True,
        "player_index": 1,
        "fields": {
            "player_name": "{player_name}",
            "school": "{school}",
            "measurables": "{height} | {weight} lbs",
            "position": "{position}",
        },
        "notes": "Same template as top target, different player.",
    },
    {
        "id": "regional_target_card",
        "name": "Regional/Storyline Target Card",
        "template_name": "HVU_Player_Feature",
        "dimensions": ["1080x1350 (Feed)"],
        "phase": "Drop +1",
        "per_player": True,
        "player_index": 2,
        "fields": {
            "player_name": "{player_name}",
            "school": "{school}",
            "measurables": "{height} | {weight} lbs",
            "position": "{position}",
        },
        "notes": "Third target — regional or storyline angle.",
    },
    {
        "id": "sweepstakes_reminder",
        "name": "Sweepstakes Reminder Card",
        "template_name": "HVU_Sweepstakes",
        "dimensions": ["1080x1080 (Square)", "1080x1920 (Story)"],
        "phase": "Drop +1",
        "fields": {
            "cta": "Earn XP. Win prizes.",
            "link": "[Platform link]",
        },
        "notes": "Sweepstakes push. Gold accent. Urgency.",
    },
    {
        "id": "keyword_hunt",
        "name": "Keyword Hunt Graphic",
        "template_name": "HVU_Keyword",
        "dimensions": ["1080x1920 (Story)"],
        "phase": "Drop +2",
        "fields": {
            "instruction": "Find the keyword in Mike V's video. +20 XP.",
            "hint": "[Keyword hint]",
        },
        "notes": "Mystery/hunt vibe. Dark background.",
    },
    {
        "id": "quiz_promo",
        "name": "Quiz Promo Card",
        "template_name": "HVU_Quiz_Promo",
        "dimensions": ["1080x1080 (Square)"],
        "phase": "Drop +1",
        "fields": {
            "quiz_title": "Which {year} {position} Are You?",
            "cta": "Take the quiz. Earn XP.",
        },
        "notes": "Personality quiz promo. Fun, engaging.",
    },
    {
        "id": "recap",
        "name": "Recap Graphic",
        "template_name": "HVU_Recap",
        "dimensions": ["1080x1080 (Square)"],
        "phase": "Drop +2",
        "fields": {
            "position_group": "{position}",
            "year": "{year}",
            "headline": "{position} DROP — THE RECAP",
        },
        "notes": "Clean wrap-up graphic. Navy. Summary feel.",
    },
]


def _resolve_fields(
    fields: dict,
    position: str,
    year: int,
    players: list[dict],
    player_index: int | None = None,
) -> dict[str, str]:
    """Replace template tokens with actual player/drop data."""
    all_names = ", ".join(p["name"] for p in players)
    redacted = ", ".join("████" for _ in players)
    resolved = {}

    player = players[player_index] if player_index is not None and player_index < len(players) else {}
    measurables = player.get("measurables", {})

    replacements = {
        "{position}": position.upper(),
        "{year}": str(year),
        "{n_players}": str(len(players)),
        "{all_names}": all_names,
        "{redacted}": redacted,
        "{player_name}": player.get("name", "TBD"),
        "{school}": player.get("school", "TBD"),
        "{height}": measurables.get("height", "TBD"),
        "{weight}": measurables.get("weight", "TBD"),
        "{city}": player.get("city", ""),
        "{state}": player.get("state", ""),
        "{top_stat}": "[Key stat — from research]",
        "{context_headline}": f"WHY {position.upper()} MATTERS NOW",
        "{context_subtext}": f"Penn State's {year} {position.upper()} board is live.",
    }

    for key, template_val in fields.items():
        val = str(template_val)
        for token, replacement in replacements.items():
            val = val.replace(token, replacement)
        resolved[key] = val

    return resolved


def build_graphics_brief(
    position: str,
    year: int,
    players: list[dict],
    drop_date: str,
    output_dir: Path,
) -> tuple[Path, Path, Path]:
    """
    Produce three graphics reference files:
    1. Graphics Brief (.docx) — visual reference for Bryan in Nano Banana
    2. Graphics Data Feed (.csv) — field values per graphic for template swaps
    3. Graphics Manifest (.json) — machine-readable for future API integration
    """
    pos = position.upper()
    output_dir.mkdir(parents=True, exist_ok=True)

    # ═══════════════════════════════════════════════════════════════════════
    # 1. GRAPHICS BRIEF (.docx)
    # ═══════════════════════════════════════════════════════════════════════
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(0.75))

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"NANO BANANA GRAPHICS BRIEF", size=Pt(16), color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, f"{year} {pos} INTEL DROP | {len(players)} Targets", size=Pt(11), color=GRAY)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, f"Drop Date: {drop_date}", size=Pt(10), color=GRAY)

    doc.add_paragraph()

    # Color reference
    p = doc.add_paragraph()
    _run(p, "BRAND COLORS: ", size=Pt(9), color=NAVY, bold=True)
    _run(p, "Navy #041E42 | White #FFFFFF | Light Blue #A2D2FF | Gold #B8860B",
         size=Pt(9), color=GRAY)

    doc.add_paragraph()

    # Each graphic
    for i, template in enumerate(GRAPHIC_TEMPLATES, 1):
        player_idx = template.get("player_index")
        fields = template.get("fields", {})
        resolved = _resolve_fields(fields, position, year, players, player_idx)

        # Graphic header
        table = doc.add_table(rows=1, cols=2)
        left = table.cell(0, 0)
        _set_cell_shading(left, "041E42")
        lp = left.paragraphs[0]
        _run(lp, f"#{i}", size=Pt(12), color=WHITE, bold=True)

        right = table.cell(0, 1)
        _set_cell_shading(right, "041E42")
        rp = right.paragraphs[0]

        name = template["name"].replace("{position}", pos)
        _run(rp, name, size=Pt(11), color=WHITE, bold=True)
        _run(rp, f"  |  {template['phase']}", size=Pt(9), color=RGBColor(0xA2, 0xD2, 0xFF))

        # Template & dimensions
        p = doc.add_paragraph()
        _run(p, "Nano Banana Template: ", size=Pt(9), bold=True, color=NAVY)
        _run(p, template.get("template_name", "N/A"), size=Pt(9), color=GOLD)

        p = doc.add_paragraph()
        _run(p, "Dimensions: ", size=Pt(9), bold=True, color=NAVY)
        _run(p, " | ".join(template["dimensions"]), size=Pt(9))

        # Fields to swap
        if resolved:
            p = doc.add_paragraph()
            _run(p, "FIELDS TO SWAP IN NANO BANANA:", size=Pt(9), bold=True, color=NAVY)
            for field_name, field_val in resolved.items():
                fp = doc.add_paragraph()
                _run(fp, f"  {field_name}: ", size=Pt(9), bold=True)
                _run(fp, field_val, size=Pt(9), color=RGBColor(0x33, 0x33, 0x33))

        # Per-card instructions (carousel)
        if "per_card" in template:
            p = doc.add_paragraph()
            _run(p, "CARD-BY-CARD:", size=Pt(9), bold=True, color=NAVY)
            for card in template["per_card"]:
                card_resolved = card.replace("{position}", pos).replace("{year}", str(year))
                cp = doc.add_paragraph()
                _run(cp, f"  {card_resolved}", size=Pt(9))

        # Notes
        if template.get("notes"):
            p = doc.add_paragraph()
            _run(p, "Notes: ", size=Pt(9), bold=True, color=GRAY)
            _run(p, template["notes"], size=Pt(9), color=GRAY, italic=True)

        doc.add_paragraph()  # spacer

    # Save count summary
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"TOTAL GRAPHICS: {len(GRAPHIC_TEMPLATES)} items across all phases",
         size=Pt(11), color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, "Open each template in Nano Banana. Swap the fields listed above. Export.",
         size=Pt(10), color=GRAY, italic=True)

    brief_path = output_dir / f"HVU_{pos}_Graphics_Brief_{year}.docx"
    doc.save(str(brief_path))

    # ═══════════════════════════════════════════════════════════════════════
    # 2. GRAPHICS DATA FEED (.csv)
    # ═══════════════════════════════════════════════════════════════════════
    csv_path = output_dir / f"HVU_{pos}_Graphics_Data_{year}.csv"
    with csv_path.open("w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Graphic #", "Graphic Name", "Template Name", "Dimensions",
            "Phase", "Field Name", "Field Value",
        ])
        for i, template in enumerate(GRAPHIC_TEMPLATES, 1):
            player_idx = template.get("player_index")
            fields = template.get("fields", {})
            resolved = _resolve_fields(fields, position, year, players, player_idx)
            name = template["name"].replace("{position}", pos)

            if resolved:
                for field_name, field_val in resolved.items():
                    writer.writerow([
                        i, name, template.get("template_name", ""),
                        " | ".join(template["dimensions"]),
                        template["phase"], field_name, field_val,
                    ])
            else:
                writer.writerow([
                    i, name, template.get("template_name", ""),
                    " | ".join(template["dimensions"]),
                    template["phase"], "", "",
                ])

    # ═══════════════════════════════════════════════════════════════════════
    # 3. GRAPHICS MANIFEST (.json)
    # ═══════════════════════════════════════════════════════════════════════
    manifest = []
    for i, template in enumerate(GRAPHIC_TEMPLATES, 1):
        player_idx = template.get("player_index")
        fields = template.get("fields", {})
        resolved = _resolve_fields(fields, position, year, players, player_idx)

        entry = {
            "graphic_number": i,
            "name": template["name"].replace("{position}", pos),
            "template_name": template.get("template_name", ""),
            "dimensions": template["dimensions"],
            "phase": template["phase"],
            "fields": resolved,
        }
        if "per_card" in template:
            entry["cards"] = [
                c.replace("{position}", pos).replace("{year}", str(year))
                for c in template["per_card"]
            ]
        if template.get("notes"):
            entry["notes"] = template["notes"]

        manifest.append(entry)

    manifest_path = output_dir / f"HVU_{pos}_Graphics_Manifest_{year}.json"
    manifest_path.write_text(json.dumps(manifest, indent=2))

    return brief_path, csv_path, manifest_path
