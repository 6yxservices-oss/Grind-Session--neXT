"""
Deliverable 1: Scouting Article Generator
-------------------------------------------
Produces the Brady Jordan insider article as a formatted .docx file.
Uses python-docx to build the exact layout specified in the HVU skill spec.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Brand Colors ──────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x36, 0x5D)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE = RGBColor(0xA2, 0xD2, 0xFF)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "1B365D"
CREAM_HEX = "F5F5F0"
LIGHT_GRAY_HEX = "F0F0F0"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _add_bottom_border(paragraph, color_hex: str = NAVY_HEX) -> None:
    """Add a bottom border line to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color_hex}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def _set_line_spacing(paragraph, spacing: float = 1.2) -> None:
    """Set line spacing as a multiplier."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing_elem = parse_xml(
        f'<w:spacing {nsdecls("w")} w:line="{int(spacing * 240)}" w:lineRule="auto"/>'
    )
    pPr.append(spacing_elem)


def _run(paragraph, text: str, size=Pt(10.5), color=None, bold=False, italic=False):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return run


def _add_header_bar(doc: Document, year: int, position: str) -> None:
    """HVU INSIDER | [YEAR] [POSITION] INTEL header on every page."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, "HVU INSIDER", size=Pt(9), color=NAVY, bold=True)
    _run(p, "  |  ", size=Pt(9), color=GRAY)
    _run(p, f"{year} {position.upper()} INTEL", size=Pt(9), color=NAVY)
    _add_bottom_border(p)


def _add_footer(doc: Document) -> None:
    """CONFIDENTIAL footer with page number."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CONFIDENTIAL — HVU INSIDER NETWORK", size=Pt(8), color=GRAY, italic=True)


def _add_title_block(doc: Document, headline: str, subtitle: str, year: int) -> None:
    """Centered title block with headline, subtitle, byline."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, headline, size=Pt(22), color=NAVY, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, subtitle, size=Pt(12), color=GRAY, italic=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    month_year = datetime.now().strftime("%B %Y")
    _run(p3, f"By Brady Jordan | HVU Insider | {month_year}", size=Pt(10), color=GRAY)

    # Spacer
    doc.add_paragraph()


def _add_opening(doc: Document, opening_text: str) -> None:
    """2-3 paragraph opening section."""
    for para_text in opening_text.strip().split("\n\n"):
        p = doc.add_paragraph()
        _set_line_spacing(p)
        _run(p, para_text.strip(), color=RGBColor(0x33, 0x33, 0x33))


def _add_player_header(doc: Document, player: dict) -> None:
    """Navy background table with player name and measurables."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f"</w:tblBorders>"
    )
    tblPr.append(borders)

    # Left cell: name + school
    left = table.cell(0, 0)
    _set_cell_shading(left, NAVY_HEX)
    p = left.paragraphs[0]
    _run(p, player["name"].upper(), size=Pt(16), color=WHITE, bold=True)
    p2 = left.add_paragraph()
    school = player.get("school", "")
    _run(p2, school, size=Pt(10), color=LIGHT_BLUE)

    # Right cell: measurables
    right = table.cell(0, 1)
    _set_cell_shading(right, NAVY_HEX)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    measurables = player.get("measurables", {})
    parts = []
    if measurables.get("height"):
        parts.append(measurables["height"])
    if measurables.get("weight"):
        parts.append(f'{measurables["weight"]} lbs')
    if measurables.get("class"):
        parts.append(f'Class of {measurables["class"]}')
    _run(p, " | ".join(parts) if parts else "Measurables TBD", size=Pt(10), color=LIGHT_BLUE)


def _add_photo_placeholder(doc: Document, player_name: str) -> None:
    """Centered photo placeholder box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, LIGHT_GRAY_HEX)

    # Set cell dimensions
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="4320" w:type="dxa"/>')
    tcPr.append(tcW)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add vertical spacing
    _run(p, "\n\n[ INSERT PLAYER PHOTO ]\n\n", size=Pt(10), color=GRAY, italic=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, player_name, size=Pt(8), color=GRAY)


def _add_section_paragraph(doc: Document, title: str, body: str) -> None:
    """A titled section paragraph (e.g., 'The Fit')."""
    p = doc.add_paragraph()
    _run(p, title, size=Pt(11), color=NAVY, bold=True)

    p2 = doc.add_paragraph()
    _set_line_spacing(p2)
    _run(p2, body.strip(), size=Pt(10.5), color=RGBColor(0x33, 0x33, 0x33))


def _add_insider_intel(doc: Document, bullets: list[str]) -> None:
    """Gold-headed insider intel bullets section."""
    p = doc.add_paragraph()
    _run(p, "INSIDER INTEL", size=Pt(11), color=GOLD, bold=True)
    _run(p, " — via the Insider Intern Squad", size=Pt(10), color=GOLD)

    for bullet in bullets:
        bp = doc.add_paragraph()
        _run(bp, f"\u25AA  {bullet}", size=Pt(10), color=RGBColor(0x33, 0x33, 0x33))


def _add_eval_box(doc: Document, player_name: str) -> None:
    """Light cream evaluation box with fill-in fields."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, CREAM_HEX)

    # Header
    p = cell.paragraphs[0]
    _run(
        p,
        f"INSIDER SCOUT EVALUATION — {player_name.upper()} — Mike V & Insider Intern",
        size=Pt(10),
        color=NAVY,
        bold=True,
    )

    fields = [
        "Overall Grade (1-10): ___________",
        "PSU Fit Score (1-10): ___________",
        "Ceiling Comp: ___________",
        "Role in Offense/Defense: ___________",
        "Commit Probability: ___________",
        "Notes (Mike V): _______________________________________________",
        "Notes (Insider Scout Intern): _______________________________________________",
    ]
    for field in fields:
        fp = cell.add_paragraph()
        _run(fp, field, size=Pt(9.5), color=RGBColor(0x44, 0x44, 0x44))


def _add_divider(doc: Document) -> None:
    """Navy divider line between sections."""
    p = doc.add_paragraph()
    _add_bottom_border(p)


def build_scouting_article(
    position: str,
    year: int,
    headline: str,
    subtitle: str,
    opening_text: str,
    players: list[dict],
    bottom_line: str,
    output_dir: Path,
) -> Path:
    """
    Build the full scouting article .docx.

    Each player dict should contain:
        name, school, city, state, measurables (dict with height/weight/class),
        the_fit, culture_fit, lifestyle, insider_intel (list of bullet strings),
        and optionally any research-populated fields.
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Default font ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10.5)

    # ── Header & Footer ──────────────────────────────────────────────────
    _add_header_bar(doc, year, position)
    _add_footer(doc)

    # ── Title Block ──────────────────────────────────────────────────────
    _add_title_block(doc, headline, subtitle, year)

    # ── Opening ──────────────────────────────────────────────────────────
    _add_opening(doc, opening_text)

    # ── Player Profiles ──────────────────────────────────────────────────
    for player in players:
        _add_divider(doc)
        _add_player_header(doc, player)
        _add_photo_placeholder(doc, player["name"])

        _add_section_paragraph(doc, "The Fit", player.get("the_fit", ""))
        _add_section_paragraph(doc, "Campbell Culture Fit", player.get("culture_fit", ""))
        _add_section_paragraph(doc, "State College Lifestyle", player.get("lifestyle", ""))
        _add_insider_intel(doc, player.get("insider_intel", []))
        _add_eval_box(doc, player["name"])

    # ── Bottom Line ──────────────────────────────────────────────────────
    _add_divider(doc)
    p = doc.add_paragraph()
    _run(p, "THE BOTTOM LINE", size=Pt(14), color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    _set_line_spacing(p2)
    _run(p2, bottom_line.strip(), size=Pt(10.5), color=RGBColor(0x33, 0x33, 0x33))

    # ── Sign-off ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "— Brady Jordan, HVU Insider —", size=Pt(10), color=GRAY, italic=True)

    # ── Save ─────────────────────────────────────────────────────────────
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"HVU_Insider_{position.upper()}_Drop_{year}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return filepath
