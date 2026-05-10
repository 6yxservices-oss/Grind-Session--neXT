#!/usr/bin/env python3
"""Build the Club MVP strategy brief as a .docx for sharing with Ryan.

Funnel-shaped narrative:
  TOP    — sweepstakes + social capture fans into Club MVP
  MIDDLE — predictions/live keep them engaged through fight week
  BOTTOM — sponsors pay + MVP sells PPV/tickets/merch direct
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "mvp-strategy" / "07-ryan-walkthrough.docx"

ACCENT = RGBColor(0xE1, 0x1D, 0x2E)
GOLD = RGBColor(0xD4, 0xAF, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
DARK = RGBColor(0x14, 0x14, 0x14)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
BLUE = RGBColor(0x38, 0xBD, 0xF8)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)


def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=DARK, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(28)
    elif level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12)
    if size:
        run.font.size = Pt(size)
    return p


def add_para(doc, text, bold=False, color=None, size=11, align=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_eyebrow(doc, text, color=ACCENT):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.color.rgb = color
    run.bold = True
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.runs[0] if p.runs else p.add_run("")
        if isinstance(item, tuple):
            bold_part, rest = item
            run.text = ""
            r1 = p.add_run(bold_part)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(rest)
            r2.font.size = Pt(11)
        else:
            run.text = item
            run.font.size = Pt(11)


def add_two_col(doc, left_title, left_items, right_title, right_items, left_color="F3F4F6", right_color="FEE2E4"):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    hdr = table.rows[0].cells
    shade(hdr[0], left_color)
    shade(hdr[1], right_color)
    for cell, title, items in [(hdr[0], left_title, left_items), (hdr[1], right_title, right_items)]:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        for item in items:
            ip = cell.add_paragraph(style="List Bullet")
            r = ip.runs[0] if ip.runs else ip.add_run("")
            r.text = item
            r.font.size = Pt(10)
    doc.add_paragraph()


def add_funnel_step(doc, step_num, color, eyebrow, title, body, fan_what=None, mvp_what=None):
    add_eyebrow(doc, eyebrow, color=color)
    add_heading(doc, f"{step_num}. {title}", level=2)
    add_para(doc, body, size=11)
    if fan_what or mvp_what:
        table = doc.add_table(rows=1, cols=2)
        fan_cell, mvp_cell = table.rows[0].cells
        shade(fan_cell, "EFF6FF")
        shade(mvp_cell, "FEF3C7")
        fan_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        mvp_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if fan_what:
            fp = fan_cell.paragraphs[0]
            fr = fp.add_run("WHAT THE FAN DOES/SEES")
            fr.bold = True
            fr.font.size = Pt(9)
            fr.font.color.rgb = BLUE
            for line in fan_what:
                lp = fan_cell.add_paragraph()
                lr = lp.add_run("• " + line)
                lr.font.size = Pt(10)
        if mvp_what:
            mp = mvp_cell.paragraphs[0]
            mr = mp.add_run("WHAT MVP CAPTURES")
            mr.bold = True
            mr.font.size = Pt(9)
            mr.font.color.rgb = AMBER
            for k, v in mvp_what:
                line = mvp_cell.add_paragraph()
                rk = line.add_run(f"{k}: ")
                rk.font.size = Pt(10)
                rk.font.color.rgb = MUTED
                rv = line.add_run(v)
                rv.font.size = Pt(10)
                rv.bold = True
        doc.add_paragraph()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Title block
    add_eyebrow(doc, "Club MVP — How It Drives Fight Promotion — Brief for Ryan Rechten")
    title = doc.add_paragraph()
    t1 = title.add_run("From a Social Media Post to a captured fan to a paying customer — ")
    t1.bold = True
    t1.font.size = Pt(22)
    t2 = title.add_run("on the site MVP already owns.")
    t2.bold = True
    t2.font.size = Pt(22)
    t2.font.color.rgb = ACCENT

    add_para(
        doc,
        "Sweepstakes capture fans at the top of the funnel. Engagement keeps them in. Sponsorship and "
        "direct sales monetize at the bottom. The whole thing runs as a layer on "
        "mostvaluablepromotions.com — not a new app, not a new site.",
        size=12, color=MUTED,
    )

    # 00 — The simple idea
    doc.add_paragraph()
    add_heading(doc, "00  The simple idea", level=1)
    add_para(
        doc,
        "Club MVP is a logged-in layer on top of the website MVP already owns. Same domain. Same brand. "
        "Same fights. New capability: capture a fan, keep them engaged, turn them into a paying "
        "customer — for sponsors and for MVP directly.",
        size=11,
    )
    add_two_col(
        doc,
        "Today — the brochure site",
        [
            "News, fights, fighters, shop",
            "Visitors come, read, leave anonymous",
            "MVP has no idea who they were",
        ],
        "+ One button: \"Join Club MVP\"",
        [
            "Same site, plus a logged-in fan layer",
            "MVP knows who came and what they did",
            "Free re-engagement channel forever",
        ],
    )

    # 01 — The funnel overview
    doc.add_page_break()
    add_heading(doc, "01  The funnel — one page", level=1)
    add_para(doc, "Three stages. Top captures attention. Middle keeps it. Bottom monetizes it.", size=11)

    table = doc.add_table(rows=3, cols=3)
    funnel_rows = [
        ("TOP — CAPTURE", "Sweepstakes capture fans from social", "TikTok / Instagram / YouTube → \"Win VIP fight experience\" → Join Club MVP", "EFF6FF"),
        ("MIDDLE — ENGAGE", "Predictions, polls, live scorecards earn entries", "Every action = more entries · behavioral data · sponsor dwell measured", "FEF3C7"),
        ("BOTTOM — MONETIZE", "Sponsors pay for the audience · MVP sells PPV, tickets, merch", "Sweepstakes inventory + extract current sponsor value + direct fan revenue", "DCFCE7"),
    ]
    for i, (label, title, sub, color) in enumerate(funnel_rows):
        c1, c2, c3 = table.rows[i].cells
        shade(c1, color)
        shade(c2, color)
        shade(c3, color)
        r = c1.paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(title)
        r2.bold = True
        r2.font.size = Pt(11)
        r3 = c3.paragraphs[0].add_run(sub)
        r3.font.size = Pt(10)
    doc.add_paragraph()

    # 02 — TOP: how fans enter
    doc.add_page_break()
    add_heading(doc, "02  How fans enter — TOP OF FUNNEL", level=1, color=BLUE)
    add_para(
        doc,
        "Sweepstakes is the hook. Social is the channel. Mostvaluablepromotions.com is the destination. "
        "Account creation is the conversion. Four steps. Less than 60 seconds.",
        size=11,
    )

    add_funnel_step(
        doc, 1, BLUE,
        "Discovery — on social",
        "Fan sees the sweepstakes on TikTok / Instagram / YouTube",
        "Fighter selfies + paid media drive traffic. The hook is a real, big prize: \"Win ringside seats + "
        "walkout access at the next MVP fight.\" Every social impression now has a chance to convert into a "
        "captured fan — not just a view.",
    )

    add_funnel_step(
        doc, 2, BLUE,
        "Land — on mostvaluablepromotions.com",
        "Fan lands on the site they already know",
        "The destination isn't a random microsite. It's mostvaluablepromotions.com — the brand fans "
        "already trust. The \"Join Club MVP\" button is right there in the nav. Zero friction.",
    )

    add_funnel_step(
        doc, 3, BLUE,
        "Convert — account creation",
        "Fan creates a Club MVP account — the data capture moment",
        "One screen. Three fields. One tap. Email, phone (SMS opt-in), favorite fighter. The fan thinks "
        "they're entering a sweepstakes — which they are. MVP just got a verified, opted-in record.",
        fan_what=[
            "Sees: \"Almost in. One step.\"",
            "Enters email + phone + favorite fighter",
            "Taps JOIN CLUB MVP",
        ],
        mvp_what=[
            ("Fan ID", "fan_0a91f7"),
            ("Email", "captured"),
            ("SMS opt-in", "YES"),
            ("Demo / geo", "M, 24, Phoenix"),
            ("Source channel", "tiktok / creator-A"),
            ("Pixel fired", "Meta + TikTok"),
        ],
    )

    add_funnel_step(
        doc, 4, BLUE,
        "Enter — the sweepstakes",
        "Entry confirmed — and the funnel begins",
        "Fan is in. Now the platform shows them how to earn more entries — which is where the in-funnel "
        "engagement begins.",
    )

    add_para(
        doc,
        "TOP-OF-FUNNEL RESULT: Every social impression that used to evaporate now has a chance to become a "
        "named, opted-in fan. This is the new top of MVP's marketing funnel for the fight — and it costs "
        "roughly $3 per captured fan vs. $11–18 to acquire the same person on Meta cold.",
        size=11, bold=True, color=BLUE,
    )

    # 03 — MIDDLE: once they're in
    doc.add_page_break()
    add_heading(doc, "03  Once they're in — MIDDLE OF FUNNEL", level=1, color=AMBER)
    add_para(
        doc,
        "Now the platform's job is to keep the fan coming back through fight week and onto fight night. "
        "Every interaction = more sweepstakes entries for the fan, more behavioral data + sponsor dwell "
        "time for MVP.",
        size=11,
    )

    add_funnel_step(
        doc, 1, AMBER,
        "Earn entries — daily reasons to return",
        "Predictions, polls, sponsor content, referrals",
        "Every day of fight week, a new entry mechanism. Predict the winner. Vote on the walkout song. "
        "Watch a 30-second sponsor clip for bonus entries. Share with a friend. The fan is now coming "
        "back daily — pre-fight buzz MVP no longer has to rent.",
        fan_what=[
            "Picks the winner (+5)",
            "Picks the round (+5)",
            "Watches sponsor clip (+3)",
            "Shares with friends (+10)",
            "Votes on walkout song (+2)",
        ],
        mvp_what=[
            ("Predictions made", "3 / 5"),
            ("Sponsor clip dwell", "28s avg"),
            ("Shares triggered", "2"),
            ("Referral signups", "1"),
            ("Engagement score", "87 / 100"),
        ],
    )

    add_funnel_step(
        doc, 2, AMBER,
        "Re-engage — push + email",
        "MVP owns the channel now — no more paying Meta",
        "By fight week, MVP isn't paying Meta to remind fans the fight is Saturday. It just sends a push "
        "notification to every Club MVP member. T-72 hours. T-24 hours. T-2 hours. Free. Every time.",
    )

    add_funnel_step(
        doc, 3, AMBER,
        "Fight night — live participation",
        "The platform becomes a second screen",
        "Live scorecards round-by-round. Push notifications timed to the broadcast. Exclusive moments only "
        "Club MVP members get. Real-time engaged-viewer list — the holy grail for sponsors who want proof "
        "of during-broadcast attention, not after-the-fact estimates.",
    )

    add_para(
        doc,
        "MIDDLE-OF-FUNNEL RESULT: Every captured fan now has a behavioral fingerprint. Top 10% become "
        "\"Superfans\" — premium sponsor inventory + first call for ticket pre-sales. Sponsor content gets "
        "MEASURABLE DWELL TIME, not impressions. Fight-week buzz scales without renting a single new ad.",
        size=11, bold=True, color=AMBER,
    )

    # 04 — BOTTOM: monetization
    doc.add_page_break()
    add_heading(doc, "04  How it makes money — BOTTOM OF FUNNEL", level=1, color=GREEN)
    add_para(
        doc,
        "Three revenue lines, all unlocked the moment the funnel starts running. Sponsorship leads — "
        "sweepstakes is the easiest inventory in the deck. Direct fan revenue and year-round monetization "
        "compound on top.",
        size=11,
    )

    # Prong 1
    add_heading(doc, "Prong 1 — Sell sweepstakes sponsorship (NEW INVENTORY)", level=2, color=ACCENT)
    add_para(
        doc,
        "Sweepstakes are the easiest piece of inventory in the deck. Brands already have a sweepstakes "
        "line in their marketing budget. They know how the format works. They've bought it from "
        "broadcasters, sportsbooks, and creators for years. Now they buy it from MVP — with the audience "
        "file attached.",
        size=11,
    )
    add_para(
        doc,
        "Why this clears procurement fast: sweepstakes sponsorship maps to an existing approval template "
        "at every CPG, sportsbook, and beverage. No new vendor category. No new legal review.",
        size=11, italic=True,
    )
    add_bullets(doc, [
        ("Title sweepstakes", " — \"Win VIP fight night, presented by [Brand]\""),
        ("Bonus-entry mechanics", " — \"Watch [Brand]'s 30-sec clip for +5 entries\""),
        ("Branded prize tiers", " — sponsor-supplied prizes (year of Celsius, $5K travel)"),
        ("Audience hand-off", " — opt-in entrants delivered to sponsor CRM post-fight"),
    ])
    add_para(
        doc,
        "Pricing: $150K–$250K presenting · $500K–$750K fight-week · $1.5M–$3M season exclusive.",
        size=11, bold=True, color=ACCENT,
    )

    # Prong 2
    doc.add_paragraph()
    add_heading(doc, "Prong 2 — Upgrade existing sponsor deals (EXTRACT CURRENT VALUE)", level=2, color=GOLD)
    add_para(
        doc,
        "Most current MVP sponsors are paying for logo exposure — corner posts, ring mats, broadcast bugs. "
        "Same deal, three years running, same flat fee. Club MVP turns that flat fee into a measured "
        "asset — and a justification for renewal at 2–3x.",
        size=11,
    )
    add_para(
        doc,
        "The renewal pitch: \"Last year you paid $X for impressions. This year, same dollar buys "
        "impressions PLUS a sweepstakes activation, an audience file in your demo, and a measurable "
        "dwell-time report. Or for 2x, you own the activation outright.\"",
        size=11, italic=True,
    )
    add_bullets(doc, [
        ("Audit current deals", " — tag what each sponsor gets vs. could get"),
        ("Activation upsell", " — first right of refusal on Club MVP activations"),
        ("Retroactive ROI proof", " — retro-fit Dropt reporting on existing deals"),
        ("Tier ladder", " — Logo → Logo + activation → Category exclusive"),
    ])
    add_para(
        doc,
        "Expected lift on renewal: +30–100% on existing sponsor revenue with zero new logos to chase.",
        size=11, bold=True, color=GOLD,
    )

    add_para(
        doc,
        "Combined: 1–2 new sweepstakes sponsors per fight ($300K–$1M) + uplift on 4–6 existing renewals "
        "(+$500K–$2M/year). Day-one impact, no platform-side risk.",
        size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Direct fan revenue
    doc.add_paragraph()
    add_heading(doc, "Plus — direct fan revenue (this is what fills seats)", level=2)
    add_two_col(
        doc,
        "Drives THIS fight",
        [
            "Ticket sales — geo-segment captured fans → SMS at T-72h",
            "PPV bundle offers — fans who picked a winner get one-tap upsell",
            "Merch conversion — fan who picks Fighter A sees Fighter A's shirt at checkout",
            "Lookalike audiences — pixel feeds Meta/TikTok at 3–5x cold conversion",
        ],
        "Drives the NEXT fight",
        [
            "Free re-launch channel — email/SMS captured list, ~$0 CAC",
            "Pre-sale priority — superfans get first crack at tickets",
            "Off-season retention — content drops, sponsor newsletter inventory",
            "Year-round monetization — VIP/loyalty memberships, premium sweepstakes",
        ],
    )
    add_para(
        doc,
        "Year-one upside on a 250K-fan list: $2–4M in addressable media + sponsor lift — on top of "
        "per-fight sponsorship revenue above.",
        size=11, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 05 — What sponsors get (Dropt proof)
    doc.add_page_break()
    add_heading(doc, "05  What sponsors actually get — the proof slide", level=1)
    add_para(
        doc,
        "Sponsors don't ask \"how much reach?\" anymore. They ask \"how many opted-in fans, in my demo?\" "
        "Below is the live Dropt output every sponsor pitch closes with.",
        size=11,
    )

    add_heading(doc, "Sellable Audience Segments", level=2, color=BLUE)
    seg_table = doc.add_table(rows=5, cols=3)
    headers = ["Segment", "Share", "Sponsor category fit"]
    for i, h in enumerate(headers):
        c = seg_table.rows[0].cells[i]
        shade(c, "0F172A")
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    seg_rows = [
        ("The Weekly Player — engages every week, social by nature", "53%", "Sportsbooks: DraftKings, FanDuel, BetMGM"),
        ("The High-Frequency Flyer — 5+ international flights/year", "12%", "Premium travel: Marriott, Delta, AmEx Plat"),
        ("The Watch Collector — buys a watch annually+", "19%", "Luxury watches: Hublot, TAG Heuer, Breitling"),
        ("The Conscious Consumer — prioritizes healthy options", "41%", "Wellness: LMNT, Liquid IV, Celsius"),
    ]
    for i, row in enumerate(seg_rows, start=1):
        for j, val in enumerate(row):
            c = seg_table.rows[i].cells[j]
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            if j == 1:
                r.bold = True
                r.font.color.rgb = BLUE
    doc.add_paragraph()

    add_heading(doc, "Age Distribution", level=2, color=BLUE)
    add_para(doc, "80.5% of audience is under 45 — prime spending years with high CLV.", size=12, bold=True)
    add_bullets(doc, [
        "18–24: 36.9% (Next Gen / Trendsetters)",
        "25–34: 25.4% (Young Professionals)",
        "35–44: 18.2% (Peak Earners)",
        "45+: 19.5% (Established)",
    ])

    add_heading(doc, "Psychographics — The Conscious Performer", level=2, color=BLUE)
    add_bullets(doc, [
        "55.2% choose water for hydration vs 25.8% energy drinks",
        "41.3% prioritize healthy eating",
        "53% maintain an active lifestyle (weekly training)",
    ])
    add_para(
        doc,
        "Unique insight: This is NOT a stereotypical \"beer & hot dog\" combat-sports crowd. They're "
        "health-conscious performers — perfect fit for wellness brands, electrolytes (LMNT, Liquid IV), "
        "and premium activewear.",
        size=11, italic=True,
    )

    # 06 — Why this is simple
    doc.add_page_break()
    add_heading(doc, "06  Why this is simple", level=1)
    add_para(
        doc,
        "MVP doesn't build, integrate, or operate anything. The product runs on top of the website MVP "
        "already owns. The only thing MVP brings is what MVP already does best.",
        size=11,
    )
    add_two_col(
        doc,
        "MVP provides (no new lift)",
        [
            "Prizes / VIP access / experiences (already in inventory)",
            "Fight-week promotion across owned channels (already running)",
            "Sponsor relationships (already exist)",
            "Fighter participation — 15-sec selfies (already organic)",
            "Live-event amplification (already happens)",
        ],
        "Club MVP handles (everything else)",
        [
            "\"Join Club MVP\" button + member accounts on existing site",
            "Sweepstakes logic, legal compliance, prize fulfillment",
            "Sponsor integrations + branded UX",
            "Data capture, CRM, lifecycle email/SMS",
            "Dropt analytics, dashboards, sponsor reporting",
        ],
    )

    # 07 — Pilot
    add_heading(doc, "07  The 14-day pilot", level=1)
    add_para(
        doc,
        "No new tech. No new vendors. No new budget. The cost to test is effectively zero. The upside is "
        "a new revenue line MVP owns forever — and a fan funnel that compounds with every fight.",
        size=11,
    )
    add_bullets(doc, [
        ("One sweepstakes live", " — VIP fight experience as the prize, \"Join Club MVP\" button on mostvaluablepromotions.com, fighter-promoted across socials"),
        ("One sponsor pilot sold", " — even at half-rate, to land the case study"),
        ("One retargeting pixel firing", " — across MVP properties + lookalike audience seeded for the next fight"),
        ("One sponsor-ready report", " — exported from Dropt within 48h post-fight"),
    ])
    add_para(
        doc,
        "Targets: 50K+ captured fans · CPF under $3 · audience file delivered to sponsor in 48h.",
        size=11, bold=True,
    )

    # Bottom line
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THE BOTTOM LINE")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GOLD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A funnel MVP owns — for the fight, and every fight after.")
    r.bold = True
    r.font.size = Pt(16)

    add_para(
        doc,
        "Top: sweepstakes + social capture fans into Club MVP on the site MVP already owns. Middle: "
        "predictions and live participation keep them engaged through fight week. Bottom: sponsors pay "
        "for the audience and MVP sells PPV, tickets, and merch direct. Every fight makes the next one "
        "cheaper to fill.",
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Ryan — pick a fight on the calendar. Send me the date and the top 3 sponsor targets. Custom "
        "activation, audience projection, and price in 5 business days.",
        size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
